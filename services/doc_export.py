"""
doc_export.py - markdown -> .docx and .pdf conversion.

Strategy:
  1. If pandoc is on PATH, use it (best fidelity, handles tables/mermaid blocks).
  2. Else fall back to python-docx for .docx and reportlab for .pdf.

Mermaid blocks are pre-rendered to PNG via the existing /docx pipeline's
mermaid_render helpers when those are available; if not, the code-fence
is left in place with a "paste into mermaid.live" hint - clunky but
non-blocking.

This module is deliberately framework-free: every entry point takes
plain strings/paths so the deliverable bundler, the FastAPI /docx route,
and ad-hoc CLI scripts can all share it.
"""
from __future__ import annotations

import logging
import re
import shutil
import subprocess
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)


def has_pandoc() -> bool:
    """Returns True iff pandoc is callable, checking the bundled
    `bin/pandoc.exe` first (matches `_resolve_pandoc()`)."""
    import os
    here = Path(__file__).resolve().parent.parent
    local = here / "bin" / ("pandoc.exe" if os.name == "nt" else "pandoc")
    if local.exists():
        return True
    return shutil.which("pandoc") is not None


# ---------------------------------------------------------------------------
# Markdown sanitisation - shared between pandoc + fallback paths
# ---------------------------------------------------------------------------


_MERMAID_FENCE_RE = re.compile(r"```mermaid\s*\n(.*?)\n```", re.DOTALL | re.IGNORECASE)


def _strip_mermaid_blocks(md: str) -> str:
    """Replace ```mermaid ...``` fences with a labelled note - the fallback
    converters can't render mermaid, and pandoc treats it as plain code
    by default. We swap the fence for a captioned note so the doc
    still flows.

    Used only by the python-docx / reportlab fallback path (when pandoc
    can't be found AND the local mermaid renderer also fails). The
    primary path now pre-renders to PNG via `_render_mermaid_to_pngs`.
    """
    def _sub(m: re.Match) -> str:
        body = m.group(1).splitlines()
        return (
            f"\n> **Diagram (Mermaid source, paste into mermaid.live to view):**\n"
            f"> ```\n> " + "\n> ".join(body[:30]) +
            ("\n> ..." if len(body) > 30 else "") + "\n> ```\n"
        )
    return _MERMAID_FENCE_RE.sub(_sub, md or "")


def _render_mermaid_to_pngs(md: str, tmp_dir: Path) -> str:
    """Pre-render every ```mermaid``` block in `md` to a PNG inside `tmp_dir`
    and replace the fence with a `![](path){ width=6in }` markdown image.

    Tries the same three-tier render chain main.py uses (mermaid.ink ->
    mmdc -> bundled Node renderer) via the shared `_render_mermaid_local`
    helper. Falls through to the legacy "labelled note" stub when all
    renderers fail so the document still flows.

    P26 (2026-05-04): added because the deliverable bundler's docx/pdf
    output silently dropped every mermaid diagram - users opened the
    Export-ZIP docx and saw a blockquoted "paste into mermaid.live"
    note where the architecture diagram should have been.
    """
    blocks = list(_MERMAID_FENCE_RE.finditer(md or ""))
    if not blocks:
        return md

    # Lazy import - avoid creating a hard dependency on main.py at module
    # load time, and degrade cleanly when called from CLI / test contexts
    # where main.py isn't importable.
    try:
        from main import _render_mermaid_local
    except Exception as exc:
        log.debug("doc_export.mermaid.no_renderer: %s", exc)
        _render_mermaid_local = None  # type: ignore[assignment]

    # Salvage every block first so the same parser-friendly source goes
    # to both the in-browser preview and the docx/pdf rendering pipeline.
    try:
        from tools.mermaid_salvage import salvage as _salvage
    except Exception:
        _salvage = lambda x: (x, [])  # noqa: E731

    out_md = md
    # Iterate in reverse so character offsets stay stable as we splice.
    for idx, m in reversed(list(enumerate(blocks, start=1))):
        raw_code = m.group(1).strip()
        try:
            cleaned, _fixes = _salvage(raw_code)
        except Exception:
            cleaned = raw_code
        png_path = tmp_dir / f"mermaid_{idx}.png"
        rendered = False
        if _render_mermaid_local is not None:
            try:
                rendered = bool(_render_mermaid_local(cleaned, str(png_path)))
            except Exception as exc:
                log.debug("doc_export.mermaid.render_failed idx=%d: %s",
                          idx, str(exc)[:200])
                rendered = False
        if rendered and png_path.exists() and png_path.stat().st_size > 0:
            # Pandoc accepts forward-slash paths on Windows too. The
            # `{ width=6in }` attribute survives the pandoc image
            # extension; the python-docx fallback strips it via the
            # main.py docx route's regex (kept for parity).
            url = str(png_path).replace("\\", "/")
            replacement = (
                f"\n\n**System Architecture Diagram {idx}**\n\n"
                f"![Diagram {idx}]({url}){{ width=6in }}\n\n"
            )
        else:
            # Fall back to the labelled-note stub so the doc still flows.
            replacement = (
                f"\n\n**System Architecture Diagram {idx}** "
                f"*(auto-renderer unavailable - source preserved below)*\n\n"
                f"```text\n{cleaned}\n```\n\n"
            )
        out_md = out_md[:m.start()] + replacement + out_md[m.end():]
    return out_md


# ---------------------------------------------------------------------------
# Pandoc path
# ---------------------------------------------------------------------------


def _pandoc_convert(md_path: Path, out_path: Path, fmt: str,
                    resource_path: Optional[Path] = None) -> bool:
    """Run pandoc.  fmt is 'docx' or 'pdf'.  Returns True on success.

    `resource_path` (optional) is added to pandoc's image-search path so
    pre-rendered mermaid PNGs in a temp directory resolve correctly.
    """
    pandoc_bin = _resolve_pandoc()
    try:
        cmd = [pandoc_bin, str(md_path), "-o", str(out_path)]
        if resource_path is not None:
            cmd += ["--resource-path", str(resource_path)]
        if fmt == "pdf":
            # PDF needs a Latex engine; pandoc will pick the default.
            # If unavailable this raises CalledProcessError and we fall back.
            cmd += ["--pdf-engine=xelatex"]
        subprocess.run(cmd, check=True, capture_output=True, timeout=60)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError,
            subprocess.TimeoutExpired) as e:
        log.debug("pandoc failed (%s) - falling back: %s", fmt, e)
        return False


def _resolve_pandoc() -> str:
    """Mirror of main.py::_resolve_pandoc - find local bin/pandoc.exe
    first, then PATH. Returns the literal "pandoc" as last resort so
    subprocess raises FileNotFoundError and the caller falls through
    to the python-docx / reportlab fallback."""
    import os
    here = Path(__file__).resolve().parent.parent
    local = here / "bin" / ("pandoc.exe" if os.name == "nt" else "pandoc")
    if local.exists():
        return str(local)
    on_path = shutil.which("pandoc")
    if on_path:
        return on_path
    return "pandoc"


# ---------------------------------------------------------------------------
# python-docx fallback
# ---------------------------------------------------------------------------


def _docx_fallback(md_text: str, out_path: Path, title: Optional[str] = None) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        log.warning("docx fallback unavailable: python-docx not installed")
        return False

    doc = Document()
    if title:
        h = doc.add_heading(title, 0)

    for line in md_text.splitlines():
        s = line.rstrip()
        if not s:
            doc.add_paragraph("")
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(s[2:], style="List Bullet")
            p.runs[0].font.size = Pt(10) if p.runs else None
        elif s.startswith("|"):
            # Render markdown tables verbatim - python-docx tables would
            # need column-count detection ahead of time; the verbatim
            # fallback is good enough for a deliverable preview.
            p = doc.add_paragraph(s)
            for r in p.runs:
                r.font.name = "Consolas"
                r.font.size = Pt(9)
        else:
            doc.add_paragraph(s)
    doc.save(str(out_path))
    return True


# ---------------------------------------------------------------------------
# reportlab PDF fallback
# ---------------------------------------------------------------------------


def _pdf_fallback(md_text: str, out_path: Path, title: Optional[str] = None) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Preformatted,
        )
    except ImportError:
        log.warning("pdf fallback unavailable: reportlab not installed")
        return False

    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    h3 = styles["Heading3"]
    body = styles["BodyText"]
    code_style = ParagraphStyle(
        "code", parent=body, fontName="Courier", fontSize=8, leading=10,
    )

    flow = []
    if title:
        flow.append(Paragraph(title, h1))
        flow.append(Spacer(1, 6 * mm))

    in_table = False
    table_buf: list[str] = []

    def _flush_table():
        nonlocal table_buf
        if table_buf:
            flow.append(Preformatted("\n".join(table_buf), code_style))
            flow.append(Spacer(1, 3 * mm))
            table_buf = []

    for line in md_text.splitlines():
        s = line.rstrip()
        if s.startswith("|"):
            in_table = True
            table_buf.append(s)
            continue
        else:
            if in_table:
                _flush_table()
                in_table = False
        if not s:
            flow.append(Spacer(1, 3 * mm))
        elif s.startswith("# "):
            flow.append(Paragraph(s[2:], h1))
        elif s.startswith("## "):
            flow.append(Paragraph(s[3:], h2))
        elif s.startswith("### "):
            flow.append(Paragraph(s[4:], h3))
        elif s.startswith(("- ", "* ")):
            flow.append(Paragraph("&bull; " + _escape_html(s[2:]), body))
        else:
            flow.append(Paragraph(_escape_html(s), body))
    _flush_table()

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
    )
    doc.build(flow)
    return True


def _escape_html(s: str) -> str:
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;"))


# ---------------------------------------------------------------------------
# Public entry points
# ---------------------------------------------------------------------------


def md_to_docx(md_path: Path, out_path: Path, *, title: Optional[str] = None) -> bool:
    """Convert one .md file to .docx. Returns True on success.

    P26 (2026-05-04): mermaid blocks are now pre-rendered to PNGs in a
    temp directory and passed to pandoc as image references with
    `--resource-path` set to the temp dir. Pre-fix, every `mermaid` fence
    was replaced with a "paste into mermaid.live" stub via
    `_strip_mermaid_blocks`, so docx/pdf inside the per-phase Export ZIP
    had NO diagrams - just blockquoted source code.
    """
    import tempfile
    raw_md = md_path.read_text(encoding="utf-8", errors="replace")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if has_pandoc():
        # Pre-render mermaid -> PNG inside a temp dir kept alive for the
        # duration of the pandoc call, then point pandoc at that dir.
        with tempfile.TemporaryDirectory(prefix="docx_mermaid_") as tmpd:
            tmp_dir = Path(tmpd)
            md_text = _render_mermaid_to_pngs(raw_md, tmp_dir)
            sanitised = tmp_dir / (md_path.stem + ".sanitised.md")
            sanitised.write_text(md_text, encoding="utf-8")
            if _pandoc_convert(sanitised, out_path, "docx",
                               resource_path=tmp_dir):
                return True
    # Fallback path can't render mermaid - strip to a captioned note so
    # the resulting docx still has something readable in place.
    return _docx_fallback(_strip_mermaid_blocks(raw_md), out_path, title=title)


def md_to_pdf(md_path: Path, out_path: Path, *, title: Optional[str] = None) -> bool:
    """Same logic as `md_to_docx` but emits PDF via pandoc/xelatex."""
    import tempfile
    raw_md = md_path.read_text(encoding="utf-8", errors="replace")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if has_pandoc():
        with tempfile.TemporaryDirectory(prefix="pdf_mermaid_") as tmpd:
            tmp_dir = Path(tmpd)
            md_text = _render_mermaid_to_pngs(raw_md, tmp_dir)
            sanitised = tmp_dir / (md_path.stem + ".sanitised.md")
            sanitised.write_text(md_text, encoding="utf-8")
            if _pandoc_convert(sanitised, out_path, "pdf",
                               resource_path=tmp_dir):
                return True
    return _pdf_fallback(_strip_mermaid_blocks(raw_md), out_path, title=title)
